from pathlib import Path

from docx import Document as DocxDocument
from pdfminer.high_level import extract_text as pdfminer_extract_text


class UnsupportedFormatError(Exception):
    pass


def extract_text(file_path: str, extension: str) -> str:
    ext = (extension or "").lower()
    if ext == "docx":
        return _extract_docx(file_path)
    if ext == "pdf":
        return _extract_pdf(file_path)
    if ext == "txt":
        return _extract_txt(file_path)
    raise UnsupportedFormatError(f"فرمت فایل پشتیبانی نمی‌شود: {extension}")


def _extract_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def _extract_pdf(file_path: str) -> str:
    return pdfminer_extract_text(file_path).strip()


def _extract_txt(file_path: str) -> str:
    return Path(file_path).read_text(encoding="utf-8", errors="replace").strip()